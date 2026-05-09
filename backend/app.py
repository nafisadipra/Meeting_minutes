import os
import io
import re
import threading
import numpy as np
from flask import Flask, jsonify, request, Response, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from pydub import AudioSegment
import markdown as md_lib
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF

_FONT_DIR = "/System/Library/Fonts/Supplemental"
_GEORGIA_R  = f"{_FONT_DIR}/Georgia.ttf"
_GEORGIA_B  = f"{_FONT_DIR}/Georgia Bold.ttf"
_GEORGIA_I  = f"{_FONT_DIR}/Georgia Italic.ttf"
_GEORGIA_BI = f"{_FONT_DIR}/Georgia Bold Italic.ttf"

from assistant import MeetingAssistant
from profiler import VoiceProfiler

GEORGIA = "Georgia"


def _set_run_font(run, bold=False, italic=False):
    run.font.name = GEORGIA
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    run._r.get_or_add_rPr()
    rPr = run._r.rPr
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), GEORGIA)
    rFonts.set(qn("w:hAnsi"), GEORGIA)
    rPr.insert(0, rFonts)


def _add_inline(para, element):
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = para.add_run(text)
                _set_run_font(run)
        elif child.name in ("strong", "b"):
            run = para.add_run(child.get_text())
            _set_run_font(run, bold=True)
        elif child.name in ("em", "i"):
            run = para.add_run(child.get_text())
            _set_run_font(run, italic=True)
        elif child.name == "code":
            run = para.add_run(child.get_text())
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            run = para.add_run(child.get_text())
            _set_run_font(run)


def _set_heading_font(para):
    for run in para.runs:
        run.font.name = GEORGIA
        run._r.get_or_add_rPr()
        rPr = run._r.rPr
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), GEORGIA)
        rFonts.set(qn("w:hAnsi"), GEORGIA)
        rPr.insert(0, rFonts)


def _normalize_markdown(content: str) -> str:
    """Ensure a blank line precedes table blocks so the parser sees them correctly.
    Only inserts a blank line when the preceding line does NOT end with | (i.e.
    we're at the start of a table, not between its rows)."""
    return re.sub(r'([^|\n])\n(\|)', r'\1\n\n\2', content)


def markdown_to_docx(content: str) -> io.BytesIO:
    html = md_lib.markdown(_normalize_markdown(content), extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    # Default style to Georgia
    normal = doc.styles["Normal"]
    normal.font.name = GEORGIA
    normal.font.size = Pt(11)

    for level in range(1, 5):
        h = doc.styles[f"Heading {level}"]
        h.font.name = GEORGIA

    for element in soup.children:
        if isinstance(element, NavigableString):
            continue

        if element.name in ("h1", "h2", "h3", "h4"):
            level = int(element.name[1])
            para = doc.add_heading(element.get_text(), level=level)
            _set_heading_font(para)

        elif element.name == "p":
            para = doc.add_paragraph()
            para.style = doc.styles["Normal"]
            _add_inline(para, element)

        elif element.name == "ul":
            for li in element.find_all("li", recursive=False):
                para = doc.add_paragraph(style="List Bullet")
                _add_inline(para, li)

        elif element.name == "ol":
            for li in element.find_all("li", recursive=False):
                para = doc.add_paragraph(style="List Number")
                _add_inline(para, li)

        elif element.name == "table":
            rows = element.find_all("tr")
            if not rows:
                continue
            cols = max(len(r.find_all(["td", "th"])) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for j, cell in enumerate(cells[:cols]):
                    tc = table.rows[i].cells[j]
                    tc.text = ""
                    para = tc.paragraphs[0]
                    run = para.add_run(cell.get_text(strip=True))
                    _set_run_font(run, bold=(cell.name == "th"))

        elif element.name == "hr":
            doc.add_paragraph("─" * 60)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _pdf_write_inline(pdf: FPDF, element, base_size: float = 11.0):
    """Render inline children of an element, handling bold/italic spans."""
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                pdf.set_font("Georgia", size=base_size)
                pdf.write(6, text)
        elif child.name in ("strong", "b"):
            pdf.set_font("Georgia", style="B", size=base_size)
            pdf.write(6, child.get_text())
            pdf.set_font("Georgia", size=base_size)
        elif child.name in ("em", "i"):
            pdf.set_font("Georgia", style="I", size=base_size)
            pdf.write(6, child.get_text())
            pdf.set_font("Georgia", size=base_size)
        elif child.name == "code":
            pdf.set_font("Courier", size=base_size - 1)
            pdf.write(6, child.get_text())
            pdf.set_font("Georgia", size=base_size)
        else:
            pdf.set_font("Georgia", size=base_size)
            pdf.write(6, child.get_text())


def markdown_to_pdf(content: str) -> io.BytesIO:
    html = md_lib.markdown(_normalize_markdown(content), extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")

    pdf = FPDF(format="A4")
    pdf.add_font("Georgia", style="",   fname=_GEORGIA_R)
    pdf.add_font("Georgia", style="B",  fname=_GEORGIA_B)
    pdf.add_font("Georgia", style="I",  fname=_GEORGIA_I)
    pdf.add_font("Georgia", style="BI", fname=_GEORGIA_BI)
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    for element in soup.children:
        if isinstance(element, NavigableString):
            continue

        if element.name == "h1":
            pdf.set_font("Georgia", style="B", size=20)
            pdf.ln(4)
            pdf.multi_cell(0, 10, element.get_text(), align="L")
            pdf.ln(3)

        elif element.name == "h2":
            pdf.set_font("Georgia", style="B", size=15)
            pdf.ln(3)
            pdf.multi_cell(0, 8, element.get_text(), align="L")
            # thin rule under h2
            y = pdf.get_y()
            pdf.set_draw_color(180, 180, 180)
            pdf.line(20, y, 190, y)
            pdf.ln(3)

        elif element.name == "h3":
            pdf.set_font("Georgia", style="B", size=12)
            pdf.ln(2)
            pdf.multi_cell(0, 7, element.get_text(), align="L")
            pdf.ln(1)

        elif element.name == "p":
            pdf.set_font("Georgia", size=11)
            _pdf_write_inline(pdf, element)
            pdf.ln(7)

        elif element.name == "ul":
            for li in element.find_all("li", recursive=False):
                pdf.set_font("Georgia", size=11)
                pdf.set_x(25)
                pdf.write(6, "• ")
                _pdf_write_inline(pdf, li)
                pdf.ln(6)
            pdf.ln(2)

        elif element.name == "ol":
            for idx, li in enumerate(element.find_all("li", recursive=False), 1):
                pdf.set_font("Georgia", size=11)
                pdf.set_x(25)
                pdf.write(6, f"{idx}. ")
                _pdf_write_inline(pdf, li)
                pdf.ln(6)
            pdf.ln(2)

        elif element.name == "table":
            rows = element.find_all("tr")
            if not rows:
                continue
            cols = max(len(r.find_all(["td", "th"])) for r in rows)
            page_w = pdf.w - 40
            col_w = page_w / cols
            pdf.ln(2)
            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                is_header = any(c.name == "th" for c in cells)
                for j, cell in enumerate(cells[:cols]):
                    style = "B" if is_header else ""
                    pdf.set_font("Georgia", style=style, size=10)
                    pdf.set_fill_color(240, 240, 240) if is_header else pdf.set_fill_color(255, 255, 255)
                    pdf.cell(col_w, 7, cell.get_text(strip=True),
                             border=1, fill=is_header, align="L")
                pdf.ln()
            pdf.ln(3)

        elif element.name == "hr":
            pdf.ln(3)
            pdf.set_draw_color(180, 180, 180)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(3)

        elif element.name == "pre":
            pdf.set_font("Courier", size=9)
            pdf.set_fill_color(245, 245, 245)
            pdf.multi_cell(0, 5, element.get_text(), border=1, fill=True, align="L")
            pdf.ln(3)

    buf = io.BytesIO(bytes(pdf.output()))
    buf.seek(0)
    return buf

load_dotenv()
app = Flask(__name__)
CORS(app)

_state_lock = threading.Lock()
current_assistant: MeetingAssistant | None = None
profiler = VoiceProfiler()
enrolled_profiles = profiler.load_all_profiles()


@app.route("/api/enroll", methods=["POST"])
def enroll_voice() -> tuple[Response, int]:
    global enrolled_profiles
    if "audio" not in request.files or "name" not in request.form:
        return jsonify({"error": "Missing audio or name"}), 400

    audio_file = request.files["audio"]
    person_name = request.form["name"].strip()
    if not person_name:
        return jsonify({"error": "Name cannot be empty"}), 400

    try:
        audio_data = io.BytesIO(audio_file.read())
        audio_segment = AudioSegment.from_file(audio_data, format="webm")
        wav_io = io.BytesIO()
        audio_segment.export(wav_io, format="wav")
        wav_io.seek(0)

        embedding = profiler.enroll_from_audio(wav_io, person_name)
        with _state_lock:
            enrolled_profiles[person_name] = embedding
        return jsonify({"status": "Profile saved!"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/start", methods=["POST"])
def start_meeting() -> tuple[Response, int]:
    global current_assistant, enrolled_profiles
    data = request.json or {}
    attendees_list = data.get("attendees", [])

    minutes_path = os.path.join(
        os.getenv("DATA_ROOT", "./data"), os.getenv("MINUTES_FILE", "minutes.txt")
    )
    if os.path.exists(minutes_path):
        os.remove(minutes_path)

    with _state_lock:
        if current_assistant is None or not current_assistant.is_recording:
            current_assistant = MeetingAssistant(
                enrolled_profiles=enrolled_profiles,
                expected_attendees=attendees_list,
            )
            current_assistant.start()
            return jsonify({"status": "Recording started", "is_recording": True}), 200
    return jsonify({"status": "Already recording", "is_recording": True}), 200


@app.route("/api/stop", methods=["POST"])
def stop_meeting() -> tuple[Response, int]:
    global current_assistant, enrolled_profiles
    with _state_lock:
        assistant = current_assistant

    if assistant and assistant.is_recording:
        try:
            discovered_centroids = assistant.stop()
        except Exception as e:
            return jsonify({"error": str(e)}), 500

        new_profiles: dict = {}
        updated_profiles: dict = {}
        with _state_lock:
            for name, vector in discovered_centroids.items():
                vec_list = np.asarray(vector).flatten().tolist()
                if name not in enrolled_profiles:
                    new_profiles[name] = vec_list
                else:
                    updated_profiles[name] = vec_list

        return (
            jsonify(
                {
                    "status": "Stopping and summarizing...",
                    "is_recording": False,
                    "new_profiles_discovered": new_profiles,
                    "updated_profiles_available": updated_profiles,
                }
            ),
            200,
        )
    return jsonify({"status": "Not recording", "is_recording": False}), 200


@app.route("/api/save_discovered_profiles", methods=["POST"])
def save_discovered_profiles() -> tuple[Response, int]:
    global enrolled_profiles
    data = request.json or {}
    profiles_to_save = data.get("profiles", {})
    for name, vector_list in profiles_to_save.items():
        embedding = np.array(vector_list)
        file_path = os.path.join(profiler.profiles_dir, f"{name}.npy")
        np.save(file_path, embedding)
        with _state_lock:
            enrolled_profiles[name] = embedding
        print(f"Saved discovered profile: {name}")
    return jsonify({"status": "Profiles saved successfully!"}), 200


@app.route("/api/clear", methods=["POST"])
def clear_meeting() -> tuple[Response, int]:
    global current_assistant
    with _state_lock:
        if current_assistant and not current_assistant.is_recording:
            current_assistant.full_transcript = []
    minutes_path = os.path.join(
        os.getenv("DATA_ROOT", "./data"), os.getenv("MINUTES_FILE", "minutes.txt")
    )
    if os.path.exists(minutes_path):
        os.remove(minutes_path)
    return jsonify({"status": "Cleared"}), 200


@app.route("/api/transcript", methods=["GET"])
def get_transcript() -> tuple[Response, int]:
    with _state_lock:
        assistant = current_assistant
    if assistant:
        return (
            jsonify(
                {
                    "is_recording": assistant.is_recording,
                    "transcript": assistant.full_transcript,
                }
            ),
            200,
        )
    return jsonify({"is_recording": False, "transcript": []}), 200


@app.route("/api/minutes", methods=["GET"])
def get_minutes() -> tuple[Response, int]:
    minutes_path = os.path.join(
        os.getenv("DATA_ROOT", "./data"), os.getenv("MINUTES_FILE", "minutes.txt")
    )
    if os.path.exists(minutes_path):
        with open(minutes_path, "r") as f:
            content = f.read()
        return jsonify({"minutes": content}), 200
    return jsonify({"minutes": None}), 200


@app.route("/api/list_profiles", methods=["GET"])
def list_profiles() -> tuple[Response, int]:
    with _state_lock:
        names = list(enrolled_profiles.keys())
    return jsonify({"enrolled_profiles": names, "count": len(names)}), 200


@app.route("/api/delete_profile", methods=["POST"])
def delete_profile() -> tuple[Response, int]:
    global enrolled_profiles
    data = request.json or {}
    name = data.get("name")

    if not name:
        return jsonify({"error": "Name is required"}), 400

    # 1. Remove from active memory
    with _state_lock:
        if name in enrolled_profiles:
            del enrolled_profiles[name]

    # 2. Delete the physical .npy file from the disk
    file_path = os.path.join(profiler.profiles_dir, f"{name}.npy")
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            return jsonify({"error": f"Failed to delete file: {str(e)}"}), 500
            
    return jsonify({"status": f"Profile '{name}' deleted successfully"}), 200

@app.route("/api/export/docx", methods=["POST"])
def export_docx() -> Response:
    data = request.json or {}
    content = data.get("content", "")
    filename = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("filename", "Meeting_Minutes"))
    buf = markdown_to_docx(content)
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{filename}.docx",
    )


@app.route("/api/export/pdf", methods=["POST"])
def export_pdf() -> Response:
    data = request.json or {}
    content = data.get("content", "")
    filename = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("filename", "Meeting_Minutes"))
    buf = markdown_to_pdf(content)
    return send_file(
        buf,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"{filename}.pdf",
    )


if __name__ == "__main__":
    port = int(os.getenv("FLASK_PORT", 5000))
    debug = os.getenv("FLASK_DEBUG", "False").lower() == "true"
    app.run(debug=debug, port=port, use_reloader=False)
